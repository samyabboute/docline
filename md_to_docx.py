from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

md_path = r"C:\Users\user\Downloads\prospeo_v17_qa\prospeo\DOCLINE_DOSSIER_COMPLET.md"
out_path = r"C:\Users\user\Downloads\prospeo_v17_qa\prospeo\DOCLINE_DOSSIER_COMPLET.docx"

with open(md_path, encoding="utf-8") as f:
    lines = f.readlines()

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.2)

# ── Color palette ─────────────────────────────────────────────────────────────
BRAND   = RGBColor(0x3B, 0x17, 0x72)
BRAND2  = RGBColor(0x5C, 0x35, 0x9C)
DARK    = RGBColor(0x0D, 0x0F, 0x14)
MUTED   = RGBColor(0x4A, 0x50, 0x68)
SUCCESS = RGBColor(0x05, 0x96, 0x69)
DANGER  = RGBColor(0xDC, 0x26, 0x26)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT   = RGBColor(0xF4, 0xF7, 0xFF)

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold  = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "3B1772")
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def parse_inline(para, text):
    """Handle **bold**, *italic*, `code` inline."""
    pattern = r"(\*\*.*?\*\*|\*.*?\*|`.*?`)"
    parts   = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            set_font(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run = para.add_run(part)
            set_font(run)

# ── Parse and render ──────────────────────────────────────────────────────────
i = 0
table_rows = []
in_code_block = False
code_lines    = []

while i < len(lines):
    raw  = lines[i].rstrip("\n")
    line = raw.strip()

    # Code block
    if line.startswith("```"):
        if not in_code_block:
            in_code_block = True
            code_lines = []
        else:
            in_code_block = False
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(6)
            for cl in code_lines:
                run = p.add_run(cl + "\n")
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        i += 1
        continue

    if in_code_block:
        code_lines.append(raw)
        i += 1
        continue

    # Table detection
    if line.startswith("|"):
        table_rows.append([c.strip() for c in line.strip("|").split("|")])
        i += 1
        continue
    else:
        if table_rows:
            # Flush table
            data = [r for r in table_rows if not all(set(c) <= set("-: ") for c in r)]
            if data:
                cols = max(len(r) for r in data)
                tbl  = doc.add_table(rows=len(data), cols=cols)
                tbl.style = "Table Grid"
                tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                for ri, row in enumerate(data):
                    for ci, cell_text in enumerate(row):
                        if ci >= cols: break
                        cell = tbl.cell(ri, ci)
                        # Strip checkbox markers
                        cell_text = re.sub(r"^[\[✅✗🔄⭐🔴🟡🟢]\s*", "", cell_text)
                        cell_text = cell_text.replace("✅", "Oui").replace("✗", "Non").replace("🔄", "En cours")
                        p = cell.paragraphs[0]
                        p.clear()
                        run = p.add_run(cell_text)
                        if ri == 0:
                            shade_cell(cell, "3B1772")
                            set_font(run, size=9, bold=True, color=WHITE)
                        else:
                            bg = "F4F7FF" if ri % 2 == 0 else "FFFFFF"
                            shade_cell(cell, bg)
                            set_font(run, size=9, color=DARK)
                doc.add_paragraph()
            table_rows = []

    # Headings
    if line.startswith("# ") and not line.startswith("## "):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(line[2:].strip())
        run.font.name  = "Calibri"
        run.font.size  = Pt(22)
        run.bold       = True
        run.font.color.rgb = BRAND
        add_horizontal_rule(doc)
    elif line.startswith("### "):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(line[4:].strip())
        run.font.name  = "Calibri"
        run.font.size  = Pt(14)
        run.bold       = True
        run.font.color.rgb = BRAND2
    elif line.startswith("## "):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(5)
        run = p.add_run(line[3:].strip())
        run.font.name  = "Calibri"
        run.font.size  = Pt(17)
        run.bold       = True
        run.font.color.rgb = BRAND
        add_horizontal_rule(doc)
    elif line.startswith("#### "):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(line[5:].strip())
        run.font.name  = "Calibri"
        run.font.size  = Pt(12)
        run.bold       = True
        run.font.color.rgb = DARK
    elif line.startswith("##### "):
        p = doc.add_paragraph()
        run = p.add_run(line[6:].strip())
        set_font(run, size=11, bold=True, color=MUTED)

    # Horizontal rule
    elif line == "---":
        add_horizontal_rule(doc)

    # Blockquote
    elif line.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(line[2:])
        set_font(run, size=10, italic=True, color=MUTED)

    # Bullet list
    elif re.match(r"^[-*] ", line) or re.match(r"^  [-*] ", line):
        indent = Cm(1.5) if line.startswith("  ") else Cm(0.5)
        text   = re.sub(r"^  ?[-*] ", "", line)
        text   = re.sub(r"^\[[ x✅]\] ", "", text)
        p      = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent  = indent
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        parse_inline(p, text)

    # Numbered list
    elif re.match(r"^\d+\. ", line):
        text = re.sub(r"^\d+\. ", "", line)
        p    = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        parse_inline(p, text)

    # Empty line
    elif line == "":
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Normal paragraph
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        parse_inline(p, line)

    i += 1

doc.save(out_path)
print("OK Saved:", out_path)
